"""Tests for the docx parser module (A1: Word 文档入库 MVP)."""

from pathlib import Path

import pytest
from docx import Document
from docx.shared import Inches

from mneme.parsers.docx_parser import BlockDraft, extract_docx, extract_docx_bytes
from mneme.parsers import get_parser


# ── Fixtures ──────────────────────────────────────────────────────────────


def _make_docx(path: Path, *, paragraphs=None, headings=None, tables=None):
    """Helper to create a .docx file for testing."""
    doc = Document()
    for level, text in (headings or []):
        doc.add_heading(text, level=level)
    for text in (paragraphs or []):
        doc.add_paragraph(text)
    for rows in (tables or []):
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        for i, row in enumerate(rows):
            for j, cell_text in enumerate(row):
                table.rows[i].cells[j].text = cell_text
    doc.save(str(path))
    return path


@pytest.fixture
def empty_docx(tmp_path):
    return _make_docx(tmp_path / "empty.docx")


@pytest.fixture
def simple_docx(tmp_path):
    return _make_docx(
        tmp_path / "simple.docx",
        headings=[(1, "文档标题")],
        paragraphs=["这是第一段正文。", "这是第二段正文。"],
    )


@pytest.fixture
def complex_docx(tmp_path):
    return _make_docx(
        tmp_path / "complex.docx",
        headings=[(1, "项目报告"), (2, "第一章"), (3, "1.1 节")],
        paragraphs=["这是第一章的正文内容。", "这是 1.1 节的详细描述。"],
        tables=[
            [["姓名", "年龄", "城市"], ["张三", "25", "北京"], ["李四", "30", "上海"]],
        ],
    )


@pytest.fixture
def chinese_docx(tmp_path):
    return _make_docx(
        tmp_path / "chinese.docx",
        headings=[(1, "知识管理系统")],
        paragraphs=[
            "Mneme 是一个个人智能资产控制平面。",
            "它统一管理资产、知识、记忆、模型/API、Agent 接入等功能。",
        ],
    )


# ── BlockDraft 基本属性 ──────────────────────────────────────────────────


def test_block_draft_dataclass():
    """BlockDraft 应该有正确的属性。"""
    draft = BlockDraft(block_order=0, content_markdown="# Title", block_type="heading")
    assert draft.block_order == 0
    assert draft.content_markdown == "# Title"
    assert draft.block_type == "heading"


# ── 空文档 ──────────────────────────────────────────────────────────────


def test_extract_empty_docx(empty_docx):
    """空 .docx 应该返回空列表。"""
    blocks = extract_docx(empty_docx)
    assert blocks == []


# ── 简单文档 ────────────────────────────────────────────────────────────


def test_extract_simple_docx(simple_docx):
    """简单文档应该提取标题和段落。"""
    blocks = extract_docx(simple_docx)
    assert len(blocks) >= 3  # 1 heading + 2 paragraphs

    # 标题块
    heading_blocks = [b for b in blocks if b.block_type == "heading"]
    assert len(heading_blocks) == 1
    assert "文档标题" in heading_blocks[0].content_markdown
    assert heading_blocks[0].content_markdown.startswith("#")

    # 段落块
    para_blocks = [b for b in blocks if b.block_type == "paragraph"]
    assert len(para_blocks) >= 2
    assert any("第一段" in b.content_markdown for b in para_blocks)
    assert any("第二段" in b.content_markdown for b in para_blocks)


def test_extract_docx_block_order_sequential(simple_docx):
    """block_order 应该从 0 开始递增。"""
    blocks = extract_docx(simple_docx)
    orders = [b.block_order for b in blocks]
    assert orders == list(range(len(blocks)))


# ── 复杂文档 ────────────────────────────────────────────────────────────


def test_extract_complex_docx(complex_docx):
    """复杂文档应该提取多级标题、段落和表格。"""
    blocks = extract_docx(complex_docx)

    # 多级标题
    heading_blocks = [b for b in blocks if b.block_type == "heading"]
    assert len(heading_blocks) >= 3
    assert any("# 项目报告" in b.content_markdown for b in heading_blocks)
    assert any("## 第一章" in b.content_markdown for b in heading_blocks)
    assert any("### 1.1 节" in b.content_markdown for b in heading_blocks)

    # 表格转 Markdown
    table_blocks = [b for b in blocks if b.block_type == "table"]
    assert len(table_blocks) >= 1
    table_md = table_blocks[0].content_markdown
    assert "姓名" in table_md
    assert "张三" in table_md
    assert "|" in table_md  # Markdown 表格格式


# ── 中文内容 ────────────────────────────────────────────────────────────


def test_extract_chinese_docx(chinese_docx):
    """中文内容应该正确提取，无乱码。"""
    blocks = extract_docx(chinese_docx)
    all_text = " ".join(b.content_markdown for b in blocks)
    assert "知识管理系统" in all_text
    assert "Mneme" in all_text
    assert "个人智能资产控制平面" in all_text


# ── 解析器注册表 ────────────────────────────────────────────────────────


def test_extract_docx_bytes(simple_docx):
    """extract_docx_bytes 应该从 bytes 提取内容。"""
    content = simple_docx.read_bytes()
    blocks = extract_docx_bytes(content)
    assert len(blocks) >= 3
    assert any("文档标题" in b.content_markdown for b in blocks)


def test_get_parser_for_docx_mime():
    """get_parser 应该为 docx MIME 类型返回 bytes 解析器。"""
    parser = get_parser("application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    assert parser is not None
    assert parser is extract_docx_bytes


def test_get_parser_for_unknown_mime():
    """get_parser 应该为未知 MIME 类型返回 None。"""
    parser = get_parser("application/pdf")
    assert parser is None


def test_get_parser_for_text_mime():
    """纯文本不需要解析器。"""
    parser = get_parser("text/plain")
    assert parser is None
